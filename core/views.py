from functools import wraps
import json
import re
from datetime import date, datetime, timedelta, time
from urllib.parse import quote
from django.contrib.auth.hashers import check_password, make_password
from django.db.models import Sum, Count
from django.http import JsonResponse, HttpResponse
from django.shortcuts import redirect, render, get_object_or_404
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.core.mail import send_mail
from django.db import connection

from .models import (
    Pacientes,
    Citas,
    Pagos,
    Doctores,
    EstadosCita,
    MotivosConsulta,
    Usuarios,
    Roles,
    TiposPago,
    NotasClinicas,
    Especialidades,
    ConfiguracionClinica
)


def sistema_login_required(view_func):
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.session.get('usuario_id'):
            return redirect('/login/')
        return view_func(request, *args, **kwargs)
    return wrapper


def rol_requerido(roles_permitidos):
    def decorator(view_func):
        @wraps(view_func)
        def wrapper(request, *args, **kwargs):
            if not request.session.get('usuario_id'):
                return redirect('/login/')

            rol = request.session.get('rol', '')

            if rol not in roles_permitidos:
                request.session.flush()
                return redirect('/login/?sin_acceso=1')

            return view_func(request, *args, **kwargs)
        return wrapper
    return decorator


def login_page(request):
    if request.session.get('usuario_id'):
        return redirect('/dashboard/')
    return render(request, 'core/login.html')


@csrf_exempt
def login_json(request):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        username = body.get('username', '').strip()
        password = body.get('password', '').strip()

        usuario = Usuarios.objects.select_related('id_rol').filter(
            username=username,
            activo=1
        ).first()

        if not usuario:
            return JsonResponse({'ok': False, 'error': 'Usuario no encontrado'}, status=401)

        if not check_password(password, usuario.password_hash):
            return JsonResponse({'ok': False, 'error': 'Contraseña incorrecta'}, status=401)

        usuario.ultimo_acceso = timezone.now()
        usuario.save()

        request.session['usuario_id'] = usuario.id_usuario
        request.session['username'] = usuario.username
        request.session['nombre_usuario'] = f"{usuario.nombres or ''} {usuario.apellidos or ''}".strip()
        request.session['rol'] = usuario.id_rol.nombre_rol if usuario.id_rol else ''

        return JsonResponse({
            'ok': True,
            'mensaje': 'Login correcto',
            'redirect': '/dashboard/'
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


def logout_view(request):
    request.session.flush()
    return redirect('/login/')


@rol_requerido(['Administrador', 'Recepción', 'Doctor'])
def dashboard_page(request):
    return render(request, 'core/dashboard.html')


@rol_requerido(['Administrador', 'Recepción', 'Doctor'])
def dashboard_json(request):
    hoy = date.today()
    hace_7_dias = hoy - timedelta(days=7)

    citas_hoy = Citas.objects.filter(fecha_cita=hoy).count()

    ingresos_semanales = Pagos.objects.filter(
        fecha_pago__date__gte=hace_7_dias,
        fecha_pago__date__lte=hoy
    ).aggregate(total=Sum('monto'))['total'] or 0

    total_citas = Citas.objects.count()
    canceladas = Citas.objects.filter(id_estado_cita__nombre_estado__iexact='Cancelada').count()
    tasa_cancelaciones = round((canceladas / total_citas) * 100, 2) if total_citas > 0 else 0

    data = {
        'citas_hoy': citas_hoy,
        'pacientes_con_deuda': 0,
        'ingresos_semanales': float(ingresos_semanales),
        'tasa_cancelaciones': tasa_cancelaciones
    }

    return JsonResponse(data, json_dumps_params={'ensure_ascii': False})


@rol_requerido(['Administrador', 'Recepción', 'Doctor'])
def pacientes_page(request):
    return render(request, 'core/pacientes.html')


@rol_requerido(['Administrador', 'Recepción', 'Doctor'])
def pacientes_json(request):
    data = list(
        Pacientes.objects.values(
            'id_paciente',
            'nombres',
            'apellidos',
            'fecha_nacimiento',
            'sexo',
            'dpi_pasaporte',
            'direccion',
            'telefono',
            'correo',
            'ocupacion',
            'contacto_emergencia_nombre',
            'contacto_emergencia_telefono',
            'activo'
        )
    )

    return JsonResponse(data, safe=False, json_dumps_params={'ensure_ascii': False})


@csrf_exempt
@rol_requerido(['Administrador', 'Recepción'])
def crear_paciente_json(request):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        nuevo_paciente = Pacientes.objects.create(
            nombres=body.get('nombres', ''),
            apellidos=body.get('apellidos', ''),
            fecha_nacimiento=body.get('fecha_nacimiento') or None,
            sexo=body.get('sexo', ''),
            dpi_pasaporte=body.get('dpi_pasaporte', ''),
            direccion=body.get('direccion', ''),
            telefono=body.get('telefono', ''),
            correo=body.get('correo', ''),
            ocupacion=body.get('ocupacion', ''),
            contacto_emergencia_nombre=body.get('contacto_emergencia_nombre', ''),
            contacto_emergencia_telefono=body.get('contacto_emergencia_telefono', ''),
            activo=1
        )

        return JsonResponse({
            'ok': True,
            'mensaje': 'Paciente creado correctamente',
            'id_paciente': nuevo_paciente.id_paciente
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@csrf_exempt
@rol_requerido(['Administrador', 'Recepción'])
def actualizar_paciente_json(request, id_paciente):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        paciente = Pacientes.objects.get(id_paciente=id_paciente)

        paciente.nombres = body.get('nombres', '')
        paciente.apellidos = body.get('apellidos', '')
        paciente.fecha_nacimiento = body.get('fecha_nacimiento') or None
        paciente.sexo = body.get('sexo', '')
        paciente.dpi_pasaporte = body.get('dpi_pasaporte', '')
        paciente.direccion = body.get('direccion', '')
        paciente.telefono = body.get('telefono', '')
        paciente.correo = body.get('correo', '')
        paciente.ocupacion = body.get('ocupacion', '')
        paciente.contacto_emergencia_nombre = body.get('contacto_emergencia_nombre', '')
        paciente.contacto_emergencia_telefono = body.get('contacto_emergencia_telefono', '')

        paciente.save()

        return JsonResponse({'ok': True, 'mensaje': 'Paciente actualizado'}, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@csrf_exempt
@rol_requerido(['Administrador', 'Recepción'])
def desactivar_paciente_json(request, id_paciente):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        paciente = Pacientes.objects.get(id_paciente=id_paciente)
        paciente.activo = 0
        paciente.save()

        return JsonResponse({'ok': True}, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)



@csrf_exempt
@rol_requerido(['Administrador', 'Recepción'])
def activar_paciente_json(request, id_paciente):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        paciente = Pacientes.objects.get(id_paciente=id_paciente)
        paciente.activo = 1
        paciente.save()

        return JsonResponse({'ok': True}, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@rol_requerido(['Administrador', 'Recepción', 'Doctor'])
def citas_page(request):
    return render(request, 'core/citas.html')


@rol_requerido(['Administrador', 'Recepción', 'Doctor'])
def citas_json(request):
    citas = Citas.objects.select_related(
        'id_paciente',
        'id_doctor__id_usuario',
        'id_estado_cita',
        'id_motivo_consulta'
    )

    eventos = []

    for c in citas:
        paciente = f"{c.id_paciente.nombres} {c.id_paciente.apellidos}" if c.id_paciente else "Paciente"

        doctor = ""
        if c.id_doctor and c.id_doctor.id_usuario:
            doctor = f"{c.id_doctor.id_usuario.nombres} {c.id_doctor.id_usuario.apellidos}"

        estado = c.id_estado_cita.nombre_estado if c.id_estado_cita else "Sin estado"
        motivo = c.id_motivo_consulta.nombre_motivo if c.id_motivo_consulta else ""

        color = '#3b82f6'
        if estado.lower() == 'confirmada':
            color = '#22c55e'
        elif estado.lower() == 'cancelada':
            color = '#ef4444'
        elif estado.lower() == 'pendiente':
            color = '#f59e0b'
        elif estado.lower() == 'no-show':
            color = '#6b7280'
        elif estado.lower() == 'completada':
            color = '#2563eb'

        start = None
        end = None

        if c.fecha_cita and c.hora_inicio:
            start = f"{c.fecha_cita}T{c.hora_inicio}"

        if c.fecha_cita and c.hora_fin:
            end = f"{c.fecha_cita}T{c.hora_fin}"

        eventos.append({
            "id": c.id_cita,
            "title": paciente,
            "start": start,
            "end": end,
            "backgroundColor": color,
            "borderColor": color,
            "extendedProps": {
                "doctor": doctor,
                "estado": estado,
                "motivo": motivo,
                "modalidad": c.modalidad or "",
                "detalle": c.razon_consulta_detalle or "",
                "observaciones": c.observaciones or "",
                "id_paciente": c.id_paciente.id_paciente if c.id_paciente else '',
                "id_doctor": c.id_doctor.id_doctor if c.id_doctor else '',
                "id_estado_cita": c.id_estado_cita.id_estado_cita if c.id_estado_cita else '',
                "id_motivo_consulta": c.id_motivo_consulta.id_motivo_consulta if c.id_motivo_consulta else '',
                "link_jitsi": f"https://meet.jit.si/ClinicaDS-{c.id_cita}" if (c.modalidad or '').lower() == 'virtual' else ''
            }
        })

    return JsonResponse(eventos, safe=False, json_dumps_params={'ensure_ascii': False})


@rol_requerido(['Administrador', 'Recepción', 'Doctor'])
def catalogos_citas_json(request):
    pacientes = list(
        Pacientes.objects.filter(activo=1).values(
            'id_paciente',
            'nombres',
            'apellidos'
        )
    )

    doctores_qs = Doctores.objects.select_related('id_usuario').filter(activo=1)
    doctores = []

    for d in doctores_qs:
        nombre = "Doctor"
        if d.id_usuario:
            nombre = f"{d.id_usuario.nombres} {d.id_usuario.apellidos}"

        doctores.append({
            'id_doctor': d.id_doctor,
            'nombre': nombre
        })

    estados = list(
        EstadosCita.objects.filter(activo=1).values(
            'id_estado_cita',
            'nombre_estado'
        )
    )

    motivos = list(
        MotivosConsulta.objects.filter(activo=1).values(
            'id_motivo_consulta',
            'nombre_motivo'
        )
    )

    return JsonResponse({
        'pacientes': pacientes,
        'doctores': doctores,
        'estados': estados,
        'motivos': motivos
    }, json_dumps_params={'ensure_ascii': False})


@csrf_exempt
@rol_requerido(['Administrador', 'Recepción'])
def crear_cita_json(request):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        paciente = Pacientes.objects.get(id_paciente=body['id_paciente'])
        doctor = Doctores.objects.get(id_doctor=body['id_doctor'])
        estado = EstadosCita.objects.get(id_estado_cita=body['id_estado_cita'])
        id_motivo = body.get('id_motivo_consulta', '')
        if str(id_motivo) == 'otro' or not id_motivo:
            motivo, _ = MotivosConsulta.objects.get_or_create(nombre_motivo='Otro')
        else:
            motivo = MotivosConsulta.objects.get(id_motivo_consulta=id_motivo)
        creado_por = Usuarios.objects.get(id_usuario=request.session.get('usuario_id'))

        try:
            from zoneinfo import ZoneInfo
            ahora_gt = datetime.now(ZoneInfo('America/Guatemala')).replace(tzinfo=None)
        except Exception:
            ahora_gt = datetime.now()

        fecha_cita_val = date.fromisoformat(body['fecha_cita'])
        hora_parts = body['hora_inicio'].split(':')
        hora_inicio_val = time(int(hora_parts[0]), int(hora_parts[1]))
        dt_cita = datetime.combine(fecha_cita_val, hora_inicio_val)

        if dt_cita < ahora_gt:
            estados_permitidos = {'cancelada', 'no asistió', 'no asistio'}
            if estado.nombre_estado.lower() not in estados_permitidos:
                return JsonResponse({
                    'ok': False,
                    'error': 'Las citas en fechas u horas pasadas solo pueden guardarse con estado Cancelada o No asistió.'
                }, status=400)

        nueva_cita = Citas.objects.create(
            id_paciente=paciente,
            id_doctor=doctor,
            id_estado_cita=estado,
            id_motivo_consulta=motivo,
            fecha_cita=body['fecha_cita'],
            hora_inicio=body['hora_inicio'],
            hora_fin=body['hora_fin'],
            modalidad=body.get('modalidad', ''),
            razon_consulta_detalle=body.get('razon_consulta_detalle', ''),
            observaciones=body.get('observaciones', ''),
            creada_por=creado_por
        )

        link_jitsi = ''
        if (body.get('modalidad') or '').lower() == 'virtual':
            link_jitsi = f"https://meet.jit.si/ClinicaDS-{nueva_cita.id_cita}"

        return JsonResponse({
            'ok': True,
            'mensaje': 'Cita creada correctamente',
            'id_cita': nueva_cita.id_cita,
            'link_jitsi': link_jitsi
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@csrf_exempt
@rol_requerido(['Administrador', 'Recepción'])
def editar_cita_json(request, id_cita):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        cita = Citas.objects.get(id_cita=id_cita)
        nuevo_estado = EstadosCita.objects.get(id_estado_cita=body['id_estado_cita'])

        try:
            from zoneinfo import ZoneInfo
            ahora_gt = datetime.now(ZoneInfo('America/Guatemala')).replace(tzinfo=None)
        except Exception:
            ahora_gt = datetime.now()

        fecha_cita_val = date.fromisoformat(body['fecha_cita'])
        hora_parts = body['hora_inicio'].split(':')
        hora_inicio_val = time(int(hora_parts[0]), int(hora_parts[1]))
        dt_cita = datetime.combine(fecha_cita_val, hora_inicio_val)

        if dt_cita < ahora_gt:
            estados_permitidos = {'cancelada', 'no asistió', 'no asistio'}
            if nuevo_estado.nombre_estado.lower() not in estados_permitidos:
                return JsonResponse({
                    'ok': False,
                    'error': 'Las citas en fechas u horas pasadas solo pueden tener estado Cancelada o No asistió.'
                }, status=400)

        cita.id_paciente = Pacientes.objects.get(id_paciente=body['id_paciente'])
        cita.id_doctor = Doctores.objects.get(id_doctor=body['id_doctor'])
        cita.id_estado_cita = nuevo_estado
        id_motivo_edit = body.get('id_motivo_consulta', '')
        if str(id_motivo_edit) == 'otro' or not id_motivo_edit:
            cita.id_motivo_consulta, _ = MotivosConsulta.objects.get_or_create(nombre_motivo='Otro')
        else:
            cita.id_motivo_consulta = MotivosConsulta.objects.get(id_motivo_consulta=id_motivo_edit)
        cita.fecha_cita = body['fecha_cita']
        cita.hora_inicio = body['hora_inicio']
        cita.hora_fin = body['hora_fin']
        cita.modalidad = body.get('modalidad', '')
        cita.razon_consulta_detalle = body.get('razon_consulta_detalle', '')
        cita.observaciones = body.get('observaciones', '')
        cita.save()

        return JsonResponse({
            'ok': True,
            'mensaje': 'Cita actualizada correctamente'
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


actualizar_cita_json = editar_cita_json


@csrf_exempt
@rol_requerido(['Administrador', 'Recepción'])
def cancelar_cita_json(request, id_cita):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        cita = Citas.objects.get(id_cita=id_cita)
        estado_cancelada = EstadosCita.objects.get(nombre_estado__iexact='Cancelada')
        cita.id_estado_cita = estado_cancelada
        cita.save()

        return JsonResponse({
            'ok': True,
            'mensaje': 'Cita cancelada correctamente'
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@rol_requerido(['Administrador', 'Recepción'])
def pagos_page(request):
    return render(request, 'core/pagos.html')


@rol_requerido(['Administrador', 'Recepción'])
def pagos_json(request):
    pagos = Pagos.objects.select_related(
        'id_paciente',
        'id_cita',
        'id_tipo_pago'
    ).order_by('-id_pago')

    data = []

    for p in pagos:
        data.append({
            'id_pago': p.id_pago,
            'paciente': f"{p.id_paciente.nombres} {p.id_paciente.apellidos}" if p.id_paciente else '',
            'id_paciente': p.id_paciente.id_paciente if p.id_paciente else None,
            'id_cita': p.id_cita.id_cita if p.id_cita else None,
            'tipo_pago': p.id_tipo_pago.nombre_tipo_pago if p.id_tipo_pago else '',
            'id_tipo_pago': p.id_tipo_pago.id_tipo_pago if p.id_tipo_pago else None,
            'monto': float(p.monto) if p.monto else 0,
            'fecha_pago': p.fecha_pago.strftime('%Y-%m-%d %H:%M') if p.fecha_pago else '',
            'referencia_pago': p.referencia_pago or '',
            'observaciones': p.observaciones or ''
        })

    return JsonResponse(data, safe=False, json_dumps_params={'ensure_ascii': False})


@rol_requerido(['Administrador', 'Recepción'])
def pagos_catalogos_json(request):
    pacientes = list(
        Pacientes.objects.filter(activo=1).values(
            'id_paciente',
            'nombres',
            'apellidos'
        )
    )

    citas_qs = Citas.objects.select_related('id_paciente').order_by('-id_cita')
    citas = []

    for c in citas_qs:
        paciente = ''
        if c.id_paciente:
            paciente = f"{c.id_paciente.nombres} {c.id_paciente.apellidos}"

        citas.append({
            'id_cita': c.id_cita,
            'texto': f"Cita {c.id_cita} - {paciente} - {c.fecha_cita}"
        })

    tipos_pago = list(
        TiposPago.objects.filter(activo=1).values(
            'id_tipo_pago',
            'nombre_tipo_pago'
        )
    )

    return JsonResponse({
        'pacientes': pacientes,
        'citas': citas,
        'tipos_pago': tipos_pago
    }, json_dumps_params={'ensure_ascii': False})


@csrf_exempt
@rol_requerido(['Administrador', 'Recepción'])
def crear_pago_json(request):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        paciente = Pacientes.objects.get(id_paciente=body['id_paciente'])
        tipo_pago = TiposPago.objects.get(id_tipo_pago=body['id_tipo_pago'])

        cita = None
        if body.get('id_cita'):
            cita = Citas.objects.get(id_cita=body['id_cita'])

        nuevo_pago = Pagos.objects.create(
            id_paciente=paciente,
            id_cita=cita,
            id_tipo_pago=tipo_pago,
            monto=body.get('monto', 0),
            fecha_pago=timezone.now(),
            referencia_pago=body.get('referencia_pago', ''),
            observaciones=body.get('observaciones', '')
        )

        return JsonResponse({
            'ok': True,
            'mensaje': 'Pago registrado correctamente',
            'id_pago': nuevo_pago.id_pago
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@csrf_exempt
@rol_requerido(['Administrador', 'Recepción'])
def actualizar_pago_json(request, id_pago):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        pago = Pagos.objects.get(id_pago=id_pago)

        if not body.get('id_paciente'):
            return JsonResponse({'ok': False, 'error': 'El paciente es obligatorio.'}, status=400)

        if not body.get('id_tipo_pago'):
            return JsonResponse({'ok': False, 'error': 'El tipo de pago es obligatorio.'}, status=400)

        pago.id_paciente = Pacientes.objects.get(id_paciente=body['id_paciente'])
        pago.id_tipo_pago = TiposPago.objects.get(id_tipo_pago=body['id_tipo_pago'])

        pago.id_cita = None
        if body.get('id_cita'):
            pago.id_cita = Citas.objects.get(id_cita=body['id_cita'])

        pago.monto = body.get('monto', 0)
        pago.referencia_pago = body.get('referencia_pago', '').strip()
        pago.observaciones = body.get('observaciones', '').strip()
        pago.save()

        return JsonResponse({
            'ok': True,
            'mensaje': 'Pago actualizado correctamente'
        }, json_dumps_params={'ensure_ascii': False})

    except Pagos.DoesNotExist:
        return JsonResponse({'ok': False, 'error': 'Pago no encontrado.'}, status=404)
    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@csrf_exempt
@rol_requerido(['Administrador', 'Recepción'])
def eliminar_pago_json(request, id_pago):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        pago = Pagos.objects.get(id_pago=id_pago)
        pago.delete()

        return JsonResponse({
            'ok': True,
            'mensaje': 'Pago eliminado correctamente'
        }, json_dumps_params={'ensure_ascii': False})

    except Pagos.DoesNotExist:
        return JsonResponse({'ok': False, 'error': 'Pago no encontrado.'}, status=404)
    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@rol_requerido(['Administrador', 'Doctor'])
def notas_page(request):
    return render(request, 'core/notas.html')


@rol_requerido(['Administrador', 'Doctor'])
def notas_json(request):
    notas = NotasClinicas.objects.select_related(
        'id_paciente',
        'id_cita',
        'id_doctor__id_usuario'
    ).order_by('-id_nota_clinica')

    data = []

    for n in notas:
        doctor = ''
        if n.id_doctor and n.id_doctor.id_usuario:
            doctor = f"{n.id_doctor.id_usuario.nombres} {n.id_doctor.id_usuario.apellidos}"

        data.append({
            'id_nota_clinica': n.id_nota_clinica,
            'id_paciente': n.id_paciente.id_paciente if n.id_paciente else None,
            'paciente': f"{n.id_paciente.nombres} {n.id_paciente.apellidos}" if n.id_paciente else '',
            'id_cita': n.id_cita.id_cita if n.id_cita else None,
            'id_doctor': n.id_doctor.id_doctor if n.id_doctor else None,
            'doctor': doctor,
            'titulo': n.titulo or '',
            'contenido': n.contenido or '',
            'recomendaciones': n.recomendaciones or '',
            'fecha_nota': n.fecha_nota.strftime('%Y-%m-%d %H:%M') if n.fecha_nota else ''
        })

    return JsonResponse(data, safe=False, json_dumps_params={'ensure_ascii': False})


@rol_requerido(['Administrador', 'Doctor'])
def notas_catalogos_json(request):
    pacientes = list(
        Pacientes.objects.filter(activo=1).values(
            'id_paciente',
            'nombres',
            'apellidos'
        )
    )

    citas_qs = Citas.objects.select_related('id_paciente').order_by('-id_cita')
    citas = []

    for c in citas_qs:
        paciente = ''
        if c.id_paciente:
            paciente = f"{c.id_paciente.nombres} {c.id_paciente.apellidos}"

        citas.append({
            'id_cita': c.id_cita,
            'id_paciente': c.id_paciente.id_paciente if c.id_paciente else None,
            'texto': f"Cita {c.id_cita} - {paciente} - {c.fecha_cita}"
        })

    doctores_qs = Doctores.objects.select_related('id_usuario').filter(activo=1)
    doctores = []

    for d in doctores_qs:
        nombre = 'Doctor'
        if d.id_usuario:
            nombre = f"{d.id_usuario.nombres} {d.id_usuario.apellidos}"

        doctores.append({
            'id_doctor': d.id_doctor,
            'nombre': nombre
        })

    return JsonResponse({
        'pacientes': pacientes,
        'citas': citas,
        'doctores': doctores
    }, json_dumps_params={'ensure_ascii': False})


@csrf_exempt
@rol_requerido(['Administrador', 'Doctor'])
def crear_nota_json(request):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        paciente = Pacientes.objects.get(id_paciente=body['id_paciente'])
        doctor = Doctores.objects.get(id_doctor=body['id_doctor'])

        cita = None
        if body.get('id_cita'):
            cita = Citas.objects.get(id_cita=body['id_cita'])

        nueva_nota = NotasClinicas.objects.create(
            id_paciente=paciente,
            id_cita=cita,
            id_doctor=doctor,
            titulo=body.get('titulo', ''),
            contenido=body.get('contenido', ''),
            recomendaciones=body.get('recomendaciones', ''),
            fecha_nota=timezone.now()
        )

        return JsonResponse({
            'ok': True,
            'mensaje': 'Nota clínica guardada correctamente',
            'id_nota_clinica': nueva_nota.id_nota_clinica
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@rol_requerido(['Administrador'])
def doctores_page(request):
    return render(request, 'core/doctores.html')


@rol_requerido(['Administrador'])
def doctores_json(request):
    doctores = Doctores.objects.select_related(
        'id_usuario',
        'id_especialidad'
    ).order_by('-id_doctor')

    data = []

    for d in doctores:
        data.append({
            'id_doctor': d.id_doctor,
            'id_usuario': d.id_usuario.id_usuario if d.id_usuario else None,
            'doctor': f"{d.id_usuario.nombres} {d.id_usuario.apellidos}" if d.id_usuario else '',
            'correo': d.id_usuario.correo if d.id_usuario else '',
            'telefono': d.id_usuario.telefono if d.id_usuario else '',
            'id_especialidad': d.id_especialidad.id_especialidad if d.id_especialidad else None,
            'especialidad': d.id_especialidad.nombre_especialidad if d.id_especialidad else '',
            'numero_colegiado': d.numero_colegiado or '',
            'duracion_cita_minutos': d.duracion_cita_minutos or 60,
            'color_agenda': d.color_agenda or '#38bdf8',
            'activo': d.activo
        })

    return JsonResponse(data, safe=False, json_dumps_params={'ensure_ascii': False})


@rol_requerido(['Administrador'])
def doctores_catalogos_json(request):
    usuarios = list(
        Usuarios.objects.filter(activo=1).values(
            'id_usuario',
            'username',
            'nombres',
            'apellidos',
            'correo',
            'telefono'
        )
    )

    especialidades = list(
        Especialidades.objects.filter(activo=1).order_by('nombre_especialidad').values(
            'id_especialidad',
            'nombre_especialidad'
        )
    )

    usuarios_ya_doctores = list(
        Doctores.objects.filter(activo=1).values_list('id_usuario_id', flat=True)
    )

    return JsonResponse({
        'usuarios': usuarios,
        'especialidades': especialidades,
        'usuarios_ya_doctores': usuarios_ya_doctores
    }, json_dumps_params={'ensure_ascii': False})


@csrf_exempt
@rol_requerido(['Administrador'])
def crear_doctor_json(request):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        usuario = Usuarios.objects.get(id_usuario=body['id_usuario'])

        especialidad_custom = (body.get('especialidad_custom') or '').strip()
        if especialidad_custom:
            especialidad, _ = Especialidades.objects.get_or_create(
                nombre_especialidad=especialidad_custom,
                defaults={'activo': 1}
            )
        else:
            especialidad = Especialidades.objects.get(id_especialidad=body['id_especialidad'])

        correo = body.get('correo', '').strip()
        if correo and Usuarios.objects.filter(correo=correo).exclude(id_usuario=usuario.id_usuario).exists():
            return JsonResponse({'ok': False, 'error': 'El correo ya está registrado en otro usuario.'}, status=400)

        if body.get('nombres') is not None:
            usuario.nombres = body.get('nombres', '').strip()
        if body.get('apellidos') is not None:
            usuario.apellidos = body.get('apellidos', '').strip()
        if body.get('correo') is not None:
            usuario.correo = correo
        if body.get('telefono') is not None:
            usuario.telefono = body.get('telefono', '').strip()
        usuario.save()

        nuevo_doctor = Doctores.objects.create(
            id_usuario=usuario,
            id_especialidad=especialidad,
            numero_colegiado=body.get('numero_colegiado', ''),
            duracion_cita_minutos=60,
            color_agenda=body.get('color_agenda', '#38bdf8'),
            activo=1
        )

        return JsonResponse({
            'ok': True,
            'mensaje': 'Doctor creado correctamente',
            'id_doctor': nuevo_doctor.id_doctor
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@csrf_exempt
@rol_requerido(['Administrador'])
def actualizar_doctor_json(request, id_doctor):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        doctor = Doctores.objects.select_related('id_usuario', 'id_especialidad').get(id_doctor=id_doctor)
        usuario = Usuarios.objects.get(id_usuario=body['id_usuario'])

        especialidad_custom = (body.get('especialidad_custom') or '').strip()
        if especialidad_custom:
            especialidad, _ = Especialidades.objects.get_or_create(
                nombre_especialidad=especialidad_custom,
                defaults={'activo': 1}
            )
        else:
            especialidad = Especialidades.objects.get(id_especialidad=body['id_especialidad'])

        correo = body.get('correo', '').strip()
        if correo and Usuarios.objects.filter(correo=correo).exclude(id_usuario=usuario.id_usuario).exists():
            return JsonResponse({'ok': False, 'error': 'El correo ya está registrado en otro usuario.'}, status=400)

        usuario.nombres = body.get('nombres', usuario.nombres or '').strip()
        usuario.apellidos = body.get('apellidos', usuario.apellidos or '').strip()
        usuario.correo = correo
        usuario.telefono = body.get('telefono', usuario.telefono or '').strip()
        usuario.save()

        doctor.id_usuario = usuario
        doctor.id_especialidad = especialidad
        doctor.numero_colegiado = body.get('numero_colegiado', '').strip()
        doctor.color_agenda = body.get('color_agenda', '#38bdf8') or '#38bdf8'
        doctor.activo = int(body.get('activo', doctor.activo if doctor.activo is not None else 1))
        if not doctor.duracion_cita_minutos:
            doctor.duracion_cita_minutos = 60
        doctor.save()

        return JsonResponse({
            'ok': True,
            'mensaje': 'Doctor actualizado correctamente'
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)

@csrf_exempt
@rol_requerido(['Administrador'])
def desactivar_doctor_json(request, id_doctor):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        doctor = Doctores.objects.get(id_doctor=id_doctor)
        doctor.activo = 0
        doctor.save()

        return JsonResponse({'ok': True}, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@csrf_exempt
@rol_requerido(['Administrador'])
def activar_doctor_json(request, id_doctor):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        doctor = Doctores.objects.get(id_doctor=id_doctor)
        doctor.activo = 1
        doctor.save()

        return JsonResponse({'ok': True}, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@rol_requerido(['Administrador'])
def usuarios_page(request):
    return render(request, 'core/usuarios.html')


@rol_requerido(['Administrador'])
def usuarios_json(request):
    usuarios = Usuarios.objects.select_related('id_rol').order_by('-id_usuario')

    data = []

    for u in usuarios:
        data.append({
            'id_usuario': u.id_usuario,
            'id_rol': u.id_rol.id_rol if u.id_rol else None,
            'rol': u.id_rol.nombre_rol if u.id_rol else '',
            'username': u.username,
            'correo': u.correo or '',
            'telefono': u.telefono or '',
            'nombres': u.nombres or '',
            'apellidos': u.apellidos or '',
            'activo': u.activo,
            'fecha_creacion': u.fecha_creacion.strftime('%Y-%m-%d %H:%M') if u.fecha_creacion else '',
            'ultimo_acceso': u.ultimo_acceso.strftime('%Y-%m-%d %H:%M') if u.ultimo_acceso else ''
        })

    return JsonResponse(data, safe=False, json_dumps_params={'ensure_ascii': False})


@rol_requerido(['Administrador'])
def usuarios_catalogos_json(request):
    roles = list(
        Roles.objects.filter(activo=1).values(
            'id_rol',
            'nombre_rol'
        )
    )

    return JsonResponse({'roles': roles}, json_dumps_params={'ensure_ascii': False})


@csrf_exempt
@rol_requerido(['Administrador'])
def crear_usuario_json(request):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        username = body.get('username', '').strip()
        password = body.get('password', '').strip()
        correo = body.get('correo', '').strip()

        if not username or not password:
            return JsonResponse({'ok': False, 'error': 'Usuario y contraseña son obligatorios.'}, status=400)

        if Usuarios.objects.filter(username=username).exists():
            return JsonResponse({'ok': False, 'error': 'El nombre de usuario ya existe.'}, status=400)

        if correo and Usuarios.objects.filter(correo=correo).exists():
            return JsonResponse({'ok': False, 'error': 'El correo ya está registrado.'}, status=400)

        rol = Roles.objects.get(id_rol=body['id_rol'])

        nuevo_usuario = Usuarios.objects.create(
            id_rol=rol,
            username=username,
            password_hash=make_password(password),
            correo=correo,
            telefono=body.get('telefono', '').strip(),
            nombres=body.get('nombres', '').strip(),
            apellidos=body.get('apellidos', '').strip(),
            activo=1,
            fecha_creacion=timezone.now()
        )

        return JsonResponse({
            'ok': True,
            'mensaje': 'Usuario creado correctamente',
            'id_usuario': nuevo_usuario.id_usuario
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@csrf_exempt
@rol_requerido(['Administrador'])
def actualizar_usuario_json(request, id_usuario):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        usuario = Usuarios.objects.get(id_usuario=id_usuario)
        rol = Roles.objects.get(id_rol=body['id_rol'])

        username = body.get('username', '').strip()
        correo = body.get('correo', '').strip()

        if Usuarios.objects.filter(username=username).exclude(id_usuario=id_usuario).exists():
            return JsonResponse({'ok': False, 'error': 'El nombre de usuario ya existe.'}, status=400)

        if correo and Usuarios.objects.filter(correo=correo).exclude(id_usuario=id_usuario).exists():
            return JsonResponse({'ok': False, 'error': 'El correo ya está registrado.'}, status=400)

        usuario.id_rol = rol
        usuario.username = username
        usuario.correo = correo
        usuario.telefono = body.get('telefono', '').strip()
        usuario.nombres = body.get('nombres', '').strip()
        usuario.apellidos = body.get('apellidos', '').strip()

        nueva_password = body.get('password', '').strip()
        if nueva_password:
            usuario.password_hash = make_password(nueva_password)

        usuario.save()

        return JsonResponse({
            'ok': True,
            'mensaje': 'Usuario actualizado correctamente'
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@csrf_exempt
@rol_requerido(['Administrador'])
def desactivar_usuario_json(request, id_usuario):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        usuario = Usuarios.objects.get(id_usuario=id_usuario)

        if usuario.id_usuario == request.session.get('usuario_id'):
            return JsonResponse({'ok': False, 'error': 'No puedes desactivar tu propio usuario.'}, status=400)

        usuario.activo = 0
        usuario.save()

        return JsonResponse({
            'ok': True,
            'mensaje': 'Usuario desactivado correctamente'
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@csrf_exempt
@rol_requerido(['Administrador'])
def activar_usuario_json(request, id_usuario):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        usuario = Usuarios.objects.get(id_usuario=id_usuario)
        usuario.activo = 1
        usuario.save()

        return JsonResponse({
            'ok': True,
            'mensaje': 'Usuario activado correctamente'
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@rol_requerido(['Administrador'])
def roles_page(request):
    return render(request, 'core/roles.html')


@rol_requerido(['Administrador'])
def roles_json(request):
    roles = Roles.objects.all().order_by('-id_rol')

    data = []

    for r in roles:
        cantidad_usuarios = Usuarios.objects.filter(id_rol=r.id_rol).count()

        data.append({
            'id_rol': r.id_rol,
            'nombre_rol': r.nombre_rol,
            'descripcion': r.descripcion or '',
            'activo': r.activo,
            'cantidad_usuarios': cantidad_usuarios
        })

    return JsonResponse(data, safe=False, json_dumps_params={'ensure_ascii': False})


@csrf_exempt
@rol_requerido(['Administrador'])
def crear_rol_json(request):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)
        nombre_rol = body.get('nombre_rol', '').strip()

        if not nombre_rol:
            return JsonResponse({'ok': False, 'error': 'El nombre del rol es obligatorio.'}, status=400)

        if Roles.objects.filter(nombre_rol=nombre_rol).exists():
            return JsonResponse({'ok': False, 'error': 'Ese rol ya existe.'}, status=400)

        nuevo_rol = Roles.objects.create(
            nombre_rol=nombre_rol,
            descripcion=body.get('descripcion', '').strip(),
            activo=1
        )

        return JsonResponse({
            'ok': True,
            'mensaje': 'Rol creado correctamente',
            'id_rol': nuevo_rol.id_rol
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@csrf_exempt
@rol_requerido(['Administrador'])
def actualizar_rol_json(request, id_rol):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        rol = Roles.objects.get(id_rol=id_rol)
        nombre_rol = body.get('nombre_rol', '').strip()

        if not nombre_rol:
            return JsonResponse({'ok': False, 'error': 'El nombre del rol es obligatorio.'}, status=400)

        if Roles.objects.filter(nombre_rol=nombre_rol).exclude(id_rol=id_rol).exists():
            return JsonResponse({'ok': False, 'error': 'Ese rol ya existe.'}, status=400)

        rol.nombre_rol = nombre_rol
        rol.descripcion = body.get('descripcion', '').strip()
        rol.save()

        return JsonResponse({
            'ok': True,
            'mensaje': 'Rol actualizado correctamente'
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@csrf_exempt
@rol_requerido(['Administrador'])
def desactivar_rol_json(request, id_rol):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        rol = Roles.objects.get(id_rol=id_rol)

        usuarios_activos = Usuarios.objects.filter(id_rol=rol, activo=1).count()

        if usuarios_activos > 0:
            return JsonResponse({
                'ok': False,
                'error': 'No puedes desactivar un rol con usuarios activos.'
            }, status=400)

        rol.activo = 0
        rol.save()

        return JsonResponse({
            'ok': True,
            'mensaje': 'Rol desactivado correctamente'
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@csrf_exempt
@rol_requerido(['Administrador'])
def activar_rol_json(request, id_rol):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        rol = Roles.objects.get(id_rol=id_rol)
        rol.activo = 1
        rol.save()

        return JsonResponse({
            'ok': True,
            'mensaje': 'Rol activado correctamente'
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)


@rol_requerido(['Administrador'])
def configuracion_page(request):
    return render(request, 'core/configuracion.html')


@rol_requerido(['Administrador'])
def configuracion_json(request):
    config = ConfiguracionClinica.objects.filter(activo=1).first()

    if not config:
        return JsonResponse({
            'ok': False,
            'error': 'No hay configuración registrada'
        }, status=404)

    data = {
        'id_configuracion': config.id_configuracion,
        'nombre_clinica': config.nombre_clinica or '',
        'slogan': config.slogan or '',
        'direccion': config.direccion or '',
        'telefono': config.telefono or '',
        'correo': config.correo or '',
        'sitio_web': config.sitio_web or '',
        'logo_url': config.logo_url or '',
        'color_primario': config.color_primario or '#38bdf8',
        'color_secundario': config.color_secundario or '#0f172a',
        'fecha_actualizacion': config.fecha_actualizacion.strftime('%Y-%m-%d %H:%M') if config.fecha_actualizacion else ''
    }

    return JsonResponse(data, json_dumps_params={'ensure_ascii': False})


@csrf_exempt
@rol_requerido(['Administrador'])
def actualizar_configuracion_json(request):
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body = json.loads(request.body)

        config = ConfiguracionClinica.objects.filter(activo=1).first()

        if not config:
            config = ConfiguracionClinica.objects.create(
                nombre_clinica='Clínica DS',
                activo=1
            )

        config.nombre_clinica = body.get('nombre_clinica', '')
        config.slogan = body.get('slogan', '')
        config.direccion = body.get('direccion', '')
        config.telefono = body.get('telefono', '')
        config.correo = body.get('correo', '')
        config.sitio_web = body.get('sitio_web', '')
        config.logo_url = body.get('logo_url', '')
        config.color_primario = body.get('color_primario', '#38bdf8')
        config.color_secundario = body.get('color_secundario', '#0f172a')

        config.save()

        return JsonResponse({
            'ok': True,
            'mensaje': 'Configuración actualizada correctamente'
        }, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=400)




#@rol_requerido(['Administrador', 'Recepción', 'Doctor'])
#def reportes_page(request):
#    return render(request, 'core/reportes.html')



@rol_requerido(['Administrador', 'Recepción', 'Doctor'])
def _datos_reporte_general():
    def fetchone_val(sql, params=None, default=0):
        try:
            with connection.cursor() as cursor:
                cursor.execute(sql, params or [])
                row = cursor.fetchone()
            if not row:
                return default
            return row[0] if row[0] is not None else default
        except Exception as e:
            print("ERROR REPORTE fetchone:", e, sql)
            return default

    def fetchall(sql, params=None):
        try:
            with connection.cursor() as cursor:
                cursor.execute(sql, params or [])
                return cursor.fetchall()
        except Exception as e:
            print("ERROR REPORTE fetchall:", e, sql)
            return []

    hoy = timezone.localdate()
    ahora = timezone.now()

    total_pacientes = int(fetchone_val("SELECT COUNT(*) FROM pacientes WHERE activo = 1"))
    total_pacientes_general = int(fetchone_val("SELECT COUNT(*) FROM pacientes"))
    total_usuarios = int(fetchone_val("SELECT COUNT(*) FROM usuarios"))
    usuarios_activos = int(fetchone_val("SELECT COUNT(*) FROM usuarios WHERE activo = 1"))
    total_doctores = int(fetchone_val("SELECT COUNT(*) FROM doctores"))
    doctores_activos = int(fetchone_val("SELECT COUNT(*) FROM doctores WHERE activo = 1"))
    total_citas = int(fetchone_val("SELECT COUNT(*) FROM citas"))
    total_pagos = int(fetchone_val("SELECT COUNT(*) FROM pagos"))

    citas_hoy = int(fetchone_val("""
        SELECT COUNT(*)
        FROM citas c
        LEFT JOIN estados_cita e ON c.id_estado_cita = e.id_estado_cita
        WHERE c.fecha_cita = CURDATE()
          AND (e.nombre_estado IS NULL OR LOWER(e.nombre_estado) NOT IN ('cancelada', 'cancelado'))
    """))

    citas_proximas_total = int(fetchone_val("""
        SELECT COUNT(*)
        FROM citas c
        LEFT JOIN estados_cita e ON c.id_estado_cita = e.id_estado_cita
        WHERE c.fecha_cita >= CURDATE()
          AND (e.nombre_estado IS NULL OR LOWER(e.nombre_estado) NOT IN ('cancelada', 'cancelado'))
    """))

    ingresos_totales = float(fetchone_val("SELECT COALESCE(SUM(monto),0) FROM pagos", default=0) or 0)
    pagos_del_mes = float(fetchone_val("""
        SELECT COALESCE(SUM(monto),0)
        FROM pagos
        WHERE fecha_pago IS NOT NULL
          AND MONTH(fecha_pago) = MONTH(CURDATE())
          AND YEAR(fecha_pago) = YEAR(CURDATE())
    """, default=0) or 0)
    ingresos_semanales = float(fetchone_val("""
        SELECT COALESCE(SUM(monto),0)
        FROM pagos
        WHERE fecha_pago IS NOT NULL
          AND fecha_pago >= DATE_SUB(NOW(), INTERVAL 7 DAY)
    """, default=0) or 0)

    promedio_pago = round((ingresos_totales / total_pagos), 2) if total_pagos else 0
    citas_futuras_pendientes_pago = int(fetchone_val("""
        SELECT COUNT(*)
        FROM citas c
        LEFT JOIN pagos p ON c.id_cita = p.id_cita
        LEFT JOIN estados_cita e ON c.id_estado_cita = e.id_estado_cita
        WHERE c.fecha_cita >= CURDATE()
          AND p.id_pago IS NULL
          AND (e.nombre_estado IS NULL OR LOWER(e.nombre_estado) NOT IN ('cancelada', 'cancelado'))
    """))
    ingresos_potenciales = round(promedio_pago * citas_futuras_pendientes_pago, 2)

    canceladas = int(fetchone_val("""
        SELECT COUNT(*)
        FROM citas c
        INNER JOIN estados_cita e ON c.id_estado_cita = e.id_estado_cita
        WHERE LOWER(e.nombre_estado) IN ('cancelada', 'cancelado')
    """))
    tasa_cancelaciones = round((canceladas / total_citas) * 100, 2) if total_citas else 0

    citas_por_estado = []
    for estado, total in fetchall("""
        SELECT COALESCE(e.nombre_estado, 'Sin estado') AS estado, COUNT(c.id_cita) AS total
        FROM citas c
        LEFT JOIN estados_cita e ON c.id_estado_cita = e.id_estado_cita
        GROUP BY estado
        ORDER BY total DESC
    """):
        citas_por_estado.append({'estado': estado, 'total': int(total or 0)})

    pacientes_frecuentes = []
    for id_paciente, nombres, apellidos, total in fetchall("""
        SELECT p.id_paciente, p.nombres, p.apellidos, COUNT(c.id_cita) AS total
        FROM citas c
        INNER JOIN pacientes p ON c.id_paciente = p.id_paciente
        GROUP BY p.id_paciente, p.nombres, p.apellidos
        ORDER BY total DESC
        LIMIT 10
    """):
        pacientes_frecuentes.append({
            'id_paciente': id_paciente,
            'paciente': f"{nombres or ''} {apellidos or ''}".strip(),
            'total_citas': int(total or 0)
        })

    cumpleanios_mes = []
    for id_paciente, nombres, apellidos, fecha_nacimiento in fetchall("""
        SELECT id_paciente, nombres, apellidos, fecha_nacimiento
        FROM pacientes
        WHERE activo = 1
          AND fecha_nacimiento IS NOT NULL
          AND MONTH(fecha_nacimiento) = MONTH(CURDATE())
        ORDER BY DAY(fecha_nacimiento), nombres
    """):
        cumpleanios_mes.append({
            'id_paciente': id_paciente,
            'paciente': f"{nombres or ''} {apellidos or ''}".strip(),
            'fecha_nacimiento': str(fecha_nacimiento) if fecha_nacimiento else ''
        })

    ingresos_por_tipo = []
    for tipo, cantidad, total in fetchall("""
        SELECT COALESCE(t.nombre_tipo_pago, 'Sin tipo') AS tipo, COUNT(p.id_pago) AS cantidad, COALESCE(SUM(p.monto),0) AS total
        FROM pagos p
        LEFT JOIN tipos_pago t ON p.id_tipo_pago = t.id_tipo_pago
        GROUP BY tipo
        ORDER BY total DESC
    """):
        ingresos_por_tipo.append({
            'tipo_pago': tipo,
            'cantidad': int(cantidad or 0),
            'total': float(total or 0)
        })

    citas_proximas = []
    for row in fetchall("""
        SELECT
            c.id_cita,
            p.nombres,
            p.apellidos,
            ud.nombres,
            ud.apellidos,
            c.fecha_cita,
            c.hora_inicio,
            c.hora_fin,
            COALESCE(e.nombre_estado, 'Sin estado') AS estado,
            COALESCE(c.modalidad, '') AS modalidad,
            COALESCE(m.nombre_motivo, '') AS motivo,
            COALESCE(c.razon_consulta_detalle, '') AS detalle
        FROM citas c
        LEFT JOIN pacientes p ON c.id_paciente = p.id_paciente
        LEFT JOIN doctores d ON c.id_doctor = d.id_doctor
        LEFT JOIN usuarios ud ON d.id_usuario = ud.id_usuario
        LEFT JOIN estados_cita e ON c.id_estado_cita = e.id_estado_cita
        LEFT JOIN motivos_consulta m ON c.id_motivo_consulta = m.id_motivo_consulta
        WHERE c.fecha_cita >= CURDATE()
          AND (e.nombre_estado IS NULL OR LOWER(e.nombre_estado) NOT IN ('cancelada', 'cancelado'))
        ORDER BY c.fecha_cita ASC, c.hora_inicio ASC
        LIMIT 30
    """):
        (id_cita, pn, pa, dn, da, fecha_cita, hora_inicio, hora_fin, estado, modalidad, motivo, detalle) = row
        citas_proximas.append({
            'id_cita': id_cita,
            'paciente': f"{pn or ''} {pa or ''}".strip(),
            'doctor': f"{dn or ''} {da or ''}".strip(),
            'fecha_cita': str(fecha_cita) if fecha_cita else '',
            'hora_inicio': str(hora_inicio)[:5] if hora_inicio else '',
            'hora_fin': str(hora_fin)[:5] if hora_fin else '',
            'estado': estado,
            'modalidad': modalidad,
            'motivo': motivo,
            'detalle': detalle
        })

    sesiones_pendientes = []
    for row in fetchall("""
        SELECT
            c.id_cita,
            p.nombres,
            p.apellidos,
            ud.nombres,
            ud.apellidos,
            c.fecha_cita,
            c.hora_inicio,
            COALESCE(e.nombre_estado, 'Sin estado') AS estado
        FROM citas c
        LEFT JOIN pagos pg ON c.id_cita = pg.id_cita
        LEFT JOIN pacientes p ON c.id_paciente = p.id_paciente
        LEFT JOIN doctores d ON c.id_doctor = d.id_doctor
        LEFT JOIN usuarios ud ON d.id_usuario = ud.id_usuario
        LEFT JOIN estados_cita e ON c.id_estado_cita = e.id_estado_cita
        WHERE pg.id_pago IS NULL
          AND (e.nombre_estado IS NULL OR LOWER(e.nombre_estado) NOT IN ('cancelada', 'cancelado'))
        ORDER BY c.fecha_cita DESC, c.hora_inicio DESC
        LIMIT 30
    """):
        (id_cita, pn, pa, dn, da, fecha_cita, hora_inicio, estado) = row
        sesiones_pendientes.append({
            'id_cita': id_cita,
            'paciente': f"{pn or ''} {pa or ''}".strip(),
            'doctor': f"{dn or ''} {da or ''}".strip(),
            'fecha_cita': str(fecha_cita) if fecha_cita else '',
            'hora_inicio': str(hora_inicio)[:5] if hora_inicio else '',
            'estado': estado
        })

    usuarios_resumen = []
    for row in fetchall("""
        SELECT u.id_usuario, u.username, u.nombres, u.apellidos, COALESCE(r.nombre_rol, '') AS rol, u.correo, u.activo
        FROM usuarios u
        LEFT JOIN roles r ON u.id_rol = r.id_rol
        ORDER BY u.id_usuario DESC
        LIMIT 50
    """):
        id_usuario, username, nombres, apellidos, rol, correo, activo = row
        usuarios_resumen.append({
            'id_usuario': id_usuario,
            'usuario': username or '',
            'nombre': f"{nombres or ''} {apellidos or ''}".strip(),
            'rol': rol,
            'correo': correo or '',
            'activo': int(activo or 0)
        })

    doctores_resumen = []
    for row in fetchall("""
        SELECT d.id_doctor, u.nombres, u.apellidos, COALESCE(e.nombre_especialidad, '') AS especialidad,
               u.correo, u.telefono, d.numero_colegiado, d.activo
        FROM doctores d
        LEFT JOIN usuarios u ON d.id_usuario = u.id_usuario
        LEFT JOIN especialidades e ON d.id_especialidad = e.id_especialidad
        ORDER BY d.id_doctor DESC
        LIMIT 50
    """):
        id_doctor, nombres, apellidos, especialidad, correo, telefono, colegiado, activo = row
        doctores_resumen.append({
            'id_doctor': id_doctor,
            'doctor': f"{nombres or ''} {apellidos or ''}".strip(),
            'especialidad': especialidad,
            'correo': correo or '',
            'telefono': telefono or '',
            'colegiado': colegiado or '',
            'activo': int(activo or 0)
        })

    return {
        'fecha_generacion': ahora.strftime('%Y-%m-%d %H:%M:%S'),
        'total_pacientes': total_pacientes,
        'total_pacientes_general': total_pacientes_general,
        'total_usuarios': total_usuarios,
        'usuarios_activos': usuarios_activos,
        'total_doctores': total_doctores,
        'doctores_activos': doctores_activos,
        'total_citas': total_citas,
        'citas_hoy': citas_hoy,
        'citas_proximas_total': citas_proximas_total,
        'total_pagos': total_pagos,
        'ingresos_totales': ingresos_totales,
        'pagos_del_mes': pagos_del_mes,
        'ingresos_semanales': ingresos_semanales,
        'ingresos_potenciales': ingresos_potenciales,
        'citas_futuras_pendientes_pago': citas_futuras_pendientes_pago,
        'promedio_pago': promedio_pago,
        'tasa_cancelaciones': tasa_cancelaciones,
        'citas_por_estado': citas_por_estado,
        'pacientes_frecuentes': pacientes_frecuentes,
        'cumpleanios_mes': cumpleanios_mes,
        'ingresos_por_tipo': ingresos_por_tipo,
        'citas_proximas': citas_proximas,
        'sesiones_pendientes': sesiones_pendientes,
        'usuarios_resumen': usuarios_resumen,
        'doctores_resumen': doctores_resumen,
    }


# @rol_requerido(['Administrador', 'Recepción', 'Doctor'])
# def reportes_json(request):
#     try:
#         return JsonResponse(_datos_reporte_general(), json_dumps_params={'ensure_ascii': False})
#     except Exception as e:
#         print('ERROR GENERAL REPORTES:', e)
#         return JsonResponse({'ok': False, 'error': str(e)}, status=500, json_dumps_params={'ensure_ascii': False})


@rol_requerido(['Administrador', 'Recepción', 'Doctor'])
def generar_word_reporte_general(request):
    data = _datos_reporte_general()

    def valor(v):
        if v is None or v == '':
            return 'No registrado'
        return str(v)

    def agregar_tabla(documento, encabezados, filas):
        tabla = documento.add_table(rows=1, cols=len(encabezados))
        tabla.style = "Table Grid"
        for i, encabezado in enumerate(encabezados):
            tabla.rows[0].cells[i].text = str(encabezado)
        for fila_data in filas:
            fila = tabla.add_row().cells
            for i, dato in enumerate(fila_data):
                fila[i].text = valor(dato)
        return tabla

    documento = Document()
    titulo = documento.add_heading("Reporte General del Sistema Clínico", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    documento.add_paragraph("Clínica DS")
    documento.add_paragraph(f"Fecha de generación: {data.get('fecha_generacion')}")
    documento.add_paragraph("")

    documento.add_heading("Resumen general", level=2)
    agregar_tabla(documento, ["Indicador", "Valor"], [
        ["Pacientes activos", data.get('total_pacientes')],
        ["Pacientes registrados", data.get('total_pacientes_general')],
        ["Usuarios totales", data.get('total_usuarios')],
        ["Usuarios activos", data.get('usuarios_activos')],
        ["Doctores totales", data.get('total_doctores')],
        ["Doctores activos", data.get('doctores_activos')],
        ["Total citas", data.get('total_citas')],
        ["Citas de hoy", data.get('citas_hoy')],
        ["Citas próximas", data.get('citas_proximas_total')],
        ["Total pagos", data.get('total_pagos')],
        ["Ingresos totales", f"Q{float(data.get('ingresos_totales', 0)):.2f}"],
        ["Ingresos del mes", f"Q{float(data.get('pagos_del_mes', 0)):.2f}"],
        ["Ingresos semanales", f"Q{float(data.get('ingresos_semanales', 0)):.2f}"],
        ["Ingresos potenciales", f"Q{float(data.get('ingresos_potenciales', 0)):.2f}"],
        ["Tasa de cancelaciones", f"{data.get('tasa_cancelaciones')}%"],
    ])

    secciones = [
        ("Citas por estado", ["Estado", "Total"], [[x.get('estado'), x.get('total')] for x in data.get('citas_por_estado', [])]),
        ("Ingresos por tipo de pago", ["Tipo", "Cantidad", "Total"], [[x.get('tipo_pago'), x.get('cantidad'), f"Q{float(x.get('total', 0)):.2f}"] for x in data.get('ingresos_por_tipo', [])]),
        ("Citas próximas", ["ID", "Paciente", "Doctor", "Fecha", "Inicio", "Fin", "Estado", "Modalidad", "Motivo"], [[x.get('id_cita'), x.get('paciente'), x.get('doctor'), x.get('fecha_cita'), x.get('hora_inicio'), x.get('hora_fin'), x.get('estado'), x.get('modalidad'), x.get('motivo')] for x in data.get('citas_proximas', [])]),
        ("Pacientes frecuentes", ["Paciente", "Total citas"], [[x.get('paciente'), x.get('total_citas')] for x in data.get('pacientes_frecuentes', [])]),
        ("Doctores", ["ID", "Doctor", "Especialidad", "Correo", "Teléfono", "Colegiado", "Activo"], [[x.get('id_doctor'), x.get('doctor'), x.get('especialidad'), x.get('correo'), x.get('telefono'), x.get('colegiado'), "Sí" if x.get('activo') else "No"] for x in data.get('doctores_resumen', [])]),
        ("Usuarios", ["ID", "Usuario", "Nombre", "Rol", "Correo", "Activo"], [[x.get('id_usuario'), x.get('usuario'), x.get('nombre'), x.get('rol'), x.get('correo'), "Sí" if x.get('activo') else "No"] for x in data.get('usuarios_resumen', [])]),
        ("Sesiones pendientes de pago", ["Cita", "Paciente", "Doctor", "Fecha", "Hora", "Estado"], [[x.get('id_cita'), x.get('paciente'), x.get('doctor'), x.get('fecha_cita'), x.get('hora_inicio'), x.get('estado')] for x in data.get('sesiones_pendientes', [])]),
    ]

    for titulo_seccion, encabezados, filas in secciones:
        documento.add_paragraph("")
        documento.add_heading(titulo_seccion, level=2)
        agregar_tabla(documento, encabezados, filas or [["Sin datos"] + [""] * (len(encabezados) - 1)])

    for parrafo in documento.paragraphs:
        for run in parrafo.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)

    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9)

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = 'attachment; filename="reporte_general_clinica.docx"'
    documento.save(response)
    return response


@rol_requerido(['Administrador', 'Recepción', 'Doctor'])
def generar_word_paciente(request, id_paciente):
    paciente = get_object_or_404(Pacientes, id_paciente=id_paciente)

    def valor(valor_original):
        if valor_original is None or valor_original == "":
            return "No registrado"
        return str(valor_original)

    def fecha(valor_fecha):
        if not valor_fecha:
            return "No registrado"
        try:
            return valor_fecha.strftime("%Y-%m-%d")
        except Exception:
            return str(valor_fecha)

    def fecha_hora(valor_fecha):
        if not valor_fecha:
            return "No registrado"
        try:
            return valor_fecha.strftime("%Y-%m-%d %H:%M")
        except Exception:
            return str(valor_fecha)

    def hora(valor_hora):
        if not valor_hora:
            return "No registrado"
        try:
            return valor_hora.strftime("%H:%M")
        except Exception:
            return str(valor_hora)

    def aplicar_fuente(documento):
        for parrafo in documento.paragraphs:
            for run in parrafo.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)

        for tabla in documento.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        for run in parrafo.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(9)

    def agregar_tabla_dos_columnas(documento, datos):
        tabla = documento.add_table(rows=0, cols=2)
        tabla.style = "Table Grid"

        for etiqueta, dato in datos:
            fila = tabla.add_row().cells
            fila[0].text = str(etiqueta)
            fila[1].text = valor(dato)

        return tabla

    def agregar_tabla_registros(documento, encabezados, registros):
        tabla = documento.add_table(rows=1, cols=len(encabezados))
        tabla.style = "Table Grid"

        for i, encabezado in enumerate(encabezados):
            tabla.rows[0].cells[i].text = str(encabezado)

        for registro in registros:
            fila = tabla.add_row().cells
            for i, dato in enumerate(registro):
                fila[i].text = valor(dato)

        return tabla

    documento = Document()

    titulo = documento.add_heading("Reporte Clínico del Paciente", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    documento.add_paragraph("Clínica Médica DS")
    documento.add_paragraph("Documento generado desde el sistema clínico.")
    documento.add_paragraph("")

    # =========================
    # DATOS DEL PACIENTE
    # =========================
    documento.add_heading("Datos del paciente", level=2)

    datos_paciente = [
        ("ID", paciente.id_paciente),
        ("Nombres", paciente.nombres),
        ("Apellidos", paciente.apellidos),
        ("Fecha de nacimiento", fecha(paciente.fecha_nacimiento)),
        ("Sexo", paciente.sexo),
        ("DPI / Pasaporte", paciente.dpi_pasaporte),
        ("Dirección", paciente.direccion),
        ("Teléfono", paciente.telefono),
        ("Correo", paciente.correo),
        ("Ocupación", paciente.ocupacion),
        ("Contacto de emergencia", paciente.contacto_emergencia_nombre),
        ("Teléfono de emergencia", paciente.contacto_emergencia_telefono),
        ("Activo", "Sí" if paciente.activo else "No"),
    ]

    agregar_tabla_dos_columnas(documento, datos_paciente)
    documento.add_paragraph("")

    # =========================
    # CITAS
    # =========================
    documento.add_heading("Historial de citas", level=2)

    citas = Citas.objects.select_related(
        "id_doctor__id_usuario",
        "id_estado_cita",
        "id_motivo_consulta"
    ).filter(
        id_paciente=paciente
    ).order_by("-fecha_cita", "-hora_inicio", "-id_cita")

    if citas.exists():
        registros_citas = []

        for cita in citas:
            doctor = "No asignado"
            if cita.id_doctor and cita.id_doctor.id_usuario:
                doctor = f"{cita.id_doctor.id_usuario.nombres} {cita.id_doctor.id_usuario.apellidos}"

            estado = cita.id_estado_cita.nombre_estado if cita.id_estado_cita else "Sin estado"
            motivo = cita.id_motivo_consulta.nombre_motivo if cita.id_motivo_consulta else "Sin motivo"

            registros_citas.append([
                cita.id_cita,
                fecha(cita.fecha_cita),
                hora(cita.hora_inicio),
                hora(cita.hora_fin),
                doctor,
                estado,
                motivo,
                cita.modalidad,
                cita.razon_consulta_detalle,
                cita.observaciones,
            ])

        agregar_tabla_registros(
            documento,
            [
                "ID cita",
                "Fecha",
                "Inicio",
                "Fin",
                "Doctor",
                "Estado",
                "Motivo",
                "Modalidad",
                "Detalle",
                "Observaciones",
            ],
            registros_citas
        )
    else:
        documento.add_paragraph("No hay citas registradas para este paciente.")

    documento.add_paragraph("")

    # =========================
    # NOTAS CLÍNICAS
    # =========================
    documento.add_heading("Notas clínicas", level=2)

    notas = NotasClinicas.objects.select_related(
        "id_cita",
        "id_doctor__id_usuario"
    ).filter(
        id_paciente=paciente
    ).order_by("-fecha_nota", "-id_nota_clinica")

    if notas.exists():
        for nota in notas:
            doctor = "No asignado"
            if nota.id_doctor and nota.id_doctor.id_usuario:
                doctor = f"{nota.id_doctor.id_usuario.nombres} {nota.id_doctor.id_usuario.apellidos}"

            documento.add_paragraph(f"Nota #{nota.id_nota_clinica}", style="List Bullet")
            agregar_tabla_dos_columnas(
                documento,
                [
                    ("Fecha", fecha_hora(nota.fecha_nota)),
                    ("Cita relacionada", nota.id_cita.id_cita if nota.id_cita else "No registrada"),
                    ("Doctor", doctor),
                    ("Título", nota.titulo),
                    ("Contenido", nota.contenido),
                    ("Recomendaciones", nota.recomendaciones),
                ]
            )
            documento.add_paragraph("")
    else:
        documento.add_paragraph("No hay notas clínicas registradas para este paciente.")

    documento.add_paragraph("")

    # =========================
    # PAGOS
    # =========================
    documento.add_heading("Pagos registrados", level=2)

    pagos = Pagos.objects.select_related(
        "id_cita",
        "id_tipo_pago"
    ).filter(
        id_paciente=paciente
    ).order_by("-fecha_pago", "-id_pago")

    if pagos.exists():
        registros_pagos = []

        for pago in pagos:
            tipo_pago = pago.id_tipo_pago.nombre_tipo_pago if pago.id_tipo_pago else "Sin tipo"

            registros_pagos.append([
                pago.id_pago,
                pago.id_cita.id_cita if pago.id_cita else "No registrada",
                tipo_pago,
                f"Q {pago.monto}",
                fecha_hora(pago.fecha_pago),
                pago.referencia_pago,
                pago.observaciones,
            ])

        agregar_tabla_registros(
            documento,
            [
                "ID pago",
                "ID cita",
                "Tipo de pago",
                "Monto",
                "Fecha de pago",
                "Referencia",
                "Observaciones",
            ],
            registros_pagos
        )
    else:
        documento.add_paragraph("No hay pagos registrados para este paciente.")

    documento.add_paragraph("")

    # =========================
    # FACTURA / COMPROBANTE
    # =========================
    documento.add_heading("Factura / Comprobante de Pago", level=2)

    cita_factura = citas.filter(
        id_estado_cita__nombre_estado='Completada'
    ).order_by('-fecha_cita', '-hora_inicio').first() or citas.order_by('-fecha_cita', '-hora_inicio').first()

    pago_factura = pagos.first()

    tabla_factura = documento.add_table(rows=0, cols=2)
    tabla_factura.style = "Table Grid"

    def fila_factura(tabla, etiqueta, dato):
        fila = tabla.add_row().cells
        fila[0].text = str(etiqueta)
        fila[1].text = str(dato) if dato not in (None, '') else ''

    fila_factura(tabla_factura, "Paciente", f"{paciente.nombres or ''} {paciente.apellidos or ''}".strip())
    fila_factura(tabla_factura, "DPI / Pasaporte", paciente.dpi_pasaporte or '')
    fila_factura(tabla_factura, "Teléfono", paciente.telefono or '')
    fila_factura(tabla_factura, "Correo", paciente.correo or '')

    if cita_factura:
        doctor_fact = ''
        if cita_factura.id_doctor and cita_factura.id_doctor.id_usuario:
            u = cita_factura.id_doctor.id_usuario
            doctor_fact = f"{u.nombres or ''} {u.apellidos or ''}".strip()
        especialidad_fact = ''
        if cita_factura.id_doctor and cita_factura.id_doctor.id_especialidad:
            especialidad_fact = cita_factura.id_doctor.id_especialidad.nombre_especialidad or ''

        fila_factura(tabla_factura, "Fecha de consulta", fecha(cita_factura.fecha_cita))
        fila_factura(tabla_factura, "Horario", f"{hora(cita_factura.hora_inicio)} - {hora(cita_factura.hora_fin)}")
        fila_factura(tabla_factura, "Doctor", doctor_fact)
        fila_factura(tabla_factura, "Especialidad", especialidad_fact)
        fila_factura(tabla_factura, "Modalidad", cita_factura.modalidad or 'Presencial')
        fila_factura(tabla_factura, "Motivo", cita_factura.id_motivo_consulta.nombre_motivo if cita_factura.id_motivo_consulta else '')

    fila_factura(tabla_factura, "Forma de pago", pago_factura.id_tipo_pago.nombre_tipo_pago if pago_factura and pago_factura.id_tipo_pago else '')
    fila_factura(tabla_factura, "Monto (Q)", f"{pago_factura.monto}" if pago_factura and pago_factura.monto is not None else '')
    fila_factura(tabla_factura, "Referencia", pago_factura.referencia_pago if pago_factura else '')
    fila_factura(tabla_factura, "Observaciones de pago", pago_factura.observaciones if pago_factura else '')
    fila_factura(tabla_factura, "Firma del paciente", '')
    fila_factura(tabla_factura, "Sello y firma del médico", '')

    documento.add_paragraph("")

    aplicar_fuente(documento)

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    nombre_archivo = f"reporte_paciente_{paciente.id_paciente}.docx"
    response["Content-Disposition"] = f'attachment; filename="{nombre_archivo}"'

    documento.save(response)
    return response



def _hora_como_time(valor):
    if isinstance(valor, time):
        return valor

    if isinstance(valor, timedelta):
        segundos = int(valor.total_seconds())
        horas = segundos // 3600
        minutos = (segundos % 3600) // 60
        segundos = segundos % 60
        return time(horas, minutos, segundos)

    if isinstance(valor, str):
        partes = valor.split(':')
        horas = int(partes[0]) if len(partes) > 0 and partes[0] else 8
        minutos = int(partes[1]) if len(partes) > 1 and partes[1] else 0
        segundos = int(partes[2]) if len(partes) > 2 and partes[2] else 0
        return time(horas, minutos, segundos)

    return time(8, 0, 0)


def _rango_recordatorios_dashboard():
    # No usamos timezone.localdate() porque en producción puede fallar si Django maneja datetime naive.
    # Para recordatorios manuales basta usar fecha local de Guatemala.
    try:
        from zoneinfo import ZoneInfo
        hoy = datetime.now(ZoneInfo('America/Guatemala')).date()
    except Exception:
        hoy = date.today()
    manana = hoy + timedelta(days=1)
    return hoy, manana


def _nombre_paciente_cita(cita):
    if not cita or not cita.id_paciente:
        return 'Paciente'
    return f"{cita.id_paciente.nombres or ''} {cita.id_paciente.apellidos or ''}".strip()


def _mensaje_cita_recordatorio(cita):
    hora_inicio = _hora_como_time(cita.hora_inicio)
    nombre = _nombre_paciente_cita(cita)
    modalidad = cita.modalidad or 'Presencial'
    motivo = cita.razon_consulta_detalle or 'Consulta médica'
    mensaje = (
        f"Hola {nombre}, le recordamos su cita médica programada para el "
        f"{cita.fecha_cita} a las {hora_inicio.strftime('%H:%M')}. "
        f"Modalidad: {modalidad}. Motivo: {motivo}. Clínica Nubnest."
    )
    if modalidad.lower() == 'virtual':
        link = f"https://meet.jit.si/ClinicaDS-{cita.id_cita}"
        mensaje += f" Enlace de su cita virtual: {link}"
    return mensaje


def _mensaje_cita_recordatorio_doctor(cita):
    hora_inicio = _hora_como_time(cita.hora_inicio)
    nombre_paciente = _nombre_paciente_cita(cita)
    modalidad = cita.modalidad or 'Presencial'
    motivo = cita.razon_consulta_detalle or 'Consulta médica'
    mensaje = (
        f"Recordatorio de cita: paciente {nombre_paciente} el {cita.fecha_cita} "
        f"a las {hora_inicio.strftime('%H:%M')}. "
        f"Modalidad: {modalidad}. Motivo: {motivo}. Clínica Nubnest."
    )
    if modalidad.lower() == 'virtual':
        link = f"https://meet.jit.si/ClinicaDS-{cita.id_cita}"
        mensaje += f" Enlace de la cita virtual: {link}"
    return mensaje


def _citas_recordatorio_qs():
    hoy, manana = _rango_recordatorios_dashboard()
    return Citas.objects.select_related(
        'id_paciente',
        'id_estado_cita',
        'id_doctor__id_usuario'
    ).filter(
        fecha_cita__gte=hoy,
        fecha_cita__lte=manana
    ).exclude(
        id_estado_cita__nombre_estado__iexact='Cancelada'
    ).exclude(
        id_estado_cita__nombre_estado__iexact='Cancelado'
    ).order_by('fecha_cita', 'hora_inicio')


@rol_requerido(['Administrador', 'Recepción'])
def enviar_recordatorios_correo(request):
    import traceback as _traceback
    try:
        enviados = 0
        errores = 0
        omitidos = 0
        detalle_errores = []

        try:
            with connection.cursor() as cursor:
                cursor.execute("""
                    SELECT id_recordatorio, destinatario, mensaje
                    FROM recordatorios
                    WHERE canal = %s
                      AND estado_envio = %s
                    ORDER BY fecha_programada ASC
                    LIMIT 50
                """, ['correo', 'pendiente'])
                recordatorios = cursor.fetchall()
        except Exception as e:
            return HttpResponse(
                f"Error consultando recordatorios de correo: {str(e)}",
                status=200,
                content_type="text/plain; charset=utf-8"
            )

        if not recordatorios:
            return HttpResponse(
                "No hay recordatorios pendientes por correo.",
                status=200,
                content_type="text/plain; charset=utf-8"
            )

        for id_recordatorio, destinatario, mensaje in recordatorios:
            if not destinatario:
                omitidos += 1
                detalle_errores.append(f"Recordatorio {id_recordatorio}: sin destinatario.")
                continue

            try:
                send_mail(
                    subject="Recordatorio de cita médica",
                    message=mensaje or "Recordatorio de cita médica.",
                    from_email=None,
                    recipient_list=[destinatario],
                    fail_silently=False,
                )

                with connection.cursor() as cursor:
                    cursor.execute("""
                        UPDATE recordatorios
                        SET estado_envio = %s
                        WHERE id_recordatorio = %s
                    """, ['enviado', id_recordatorio])

                enviados += 1

            except Exception as e:
                errores += 1
                detalle = f"Recordatorio {id_recordatorio} hacia {destinatario}: {str(e)}"
                detalle_errores.append(detalle)
                print("ERROR AL ENVIAR CORREO:", detalle)

        respuesta = f"Recordatorios procesados. Enviados: {enviados}. Omitidos: {omitidos}. Errores: {errores}."

        if detalle_errores:
            respuesta += "\n\nDetalle:\n" + "\n".join(detalle_errores[:10])

        return HttpResponse(
            respuesta,
            status=200,
            content_type="text/plain; charset=utf-8"
        )

    except Exception as e:
        return HttpResponse(
            f"Error inesperado al enviar correos: {str(e)}\n\n{_traceback.format_exc()}",
            status=200,
            content_type="text/plain; charset=utf-8"
        )


def limpiar_numero_whatsapp(numero):
    numero = str(numero or '')
    numero = re.sub(r'\D', '', numero)
    if len(numero) == 8:
        numero = '502' + numero
    return numero


@rol_requerido(['Administrador', 'Recepción'])
def listar_recordatorios_whatsapp(request):
    """
    Muestra WhatsApp para citas de hoy y mañana.
    Si ya existen recordatorios pendientes en la tabla recordatorios, los muestra.
    Si no existen, igual muestra enlaces manuales directos desde las citas próximas.
    """
    pendientes = []
    try:
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT id_recordatorio, destinatario, mensaje, fecha_programada
                FROM recordatorios
                WHERE canal = %s
                  AND estado_envio = %s
                ORDER BY fecha_programada ASC
            """, ['whatsapp', 'pendiente'])
            pendientes = cursor.fetchall()
    except Exception as e:
        print('ERROR CONSULTANDO WHATSAPP PENDIENTES:', e)
        pendientes = []

    citas = list(_citas_recordatorio_qs())

    html = """
    <!DOCTYPE html>
    <html lang="es">
    <head>
        <meta charset="UTF-8">
        <title>Recordatorios WhatsApp</title>
        <style>
            body { font-family: Arial, sans-serif; background:#08111f; color:#ddeeff; padding:24px; }
            h1 { color:#ddeeff; font-size:24px; margin-bottom:8px; }
            .sub { color:#7fa0c3; margin-bottom:18px; }
            .card { background:#0f1e36; border:1px solid rgba(45,212,191,.22); border-radius:14px; padding:16px; margin-bottom:14px; }
            .numero { font-weight:bold; margin-bottom:6px; }
            .fecha { color:#93b5d9; font-size:14px; margin-bottom:8px; }
            .mensaje { background:#0a1628; padding:12px; border-radius:10px; margin-bottom:12px; white-space:pre-wrap; line-height:1.5; }
            .btn { display:inline-block; background:#22c55e; color:#04110a; text-decoration:none; padding:10px 14px; border-radius:10px; font-weight:bold; margin-right:8px; }
            .btn.secondary { background:#38bdf8; color:#06111f; }
            .empty { background:#0f1e36; border:1px solid rgba(45,212,191,.22); padding:18px; border-radius:14px; color:#cfe3ff; }
            hr { border:0; border-top:1px solid rgba(45,212,191,.18); margin:18px 0; }
        </style>
    </head>
    <body>
        <h1>Recordatorios WhatsApp</h1>
        <div class="sub">Citas de hoy y mañana. Puedes abrir WhatsApp manualmente desde cada botón.</div>
    """

    total_mostrados = 0

    if pendientes:
        html += "<h3>Pendientes generados</h3>"
        for id_recordatorio, destinatario, mensaje, fecha_programada in pendientes:
            numero_limpio = limpiar_numero_whatsapp(destinatario)
            wa_url = f"https://wa.me/{numero_limpio}?text={quote(str(mensaje))}"
            html += f"""
            <div class="card">
                <div class="numero">Número: {destinatario}</div>
                <div class="fecha">Fecha programada: {fecha_programada}</div>
                <div class="mensaje">{mensaje}</div>
                <a class="btn" target="_blank" href="{wa_url}">Abrir WhatsApp</a>
                <a class="btn secondary" target="_blank" href="/recordatorios/whatsapp/{id_recordatorio}/abrir/">Abrir y marcar enviado</a>
            </div>
            """
            total_mostrados += 1

    html += "<h3>Citas de hoy y mañana</h3>"
    for cita in citas:
        if not cita.id_paciente or not cita.id_paciente.telefono:
            continue
        mensaje = _mensaje_cita_recordatorio(cita)
        numero = cita.id_paciente.telefono
        numero_limpio = limpiar_numero_whatsapp(numero)
        wa_url = f"https://wa.me/{numero_limpio}?text={quote(str(mensaje))}"
        hora_inicio = _hora_como_time(cita.hora_inicio)
        html += f"""
        <div class="card">
            <div class="numero">Paciente: {_nombre_paciente_cita(cita)} | Número: {numero}</div>
            <div class="fecha">Cita: {cita.fecha_cita} {hora_inicio.strftime('%H:%M')}</div>
            <div class="mensaje">{mensaje}</div>
            <a class="btn" target="_blank" href="{wa_url}">Enviar WhatsApp manual</a>
            <a class="btn secondary" target="_blank" href="/recordatorios/whatsapp/cita/{cita.id_cita}/abrir/">Abrir y registrar</a>
        </div>
        """
        total_mostrados += 1

    if total_mostrados == 0:
        html += '<div class="empty">No hay citas de hoy o mañana con teléfono registrado.</div>'

    html += """
    </body>
    </html>
    """

    return HttpResponse(html)


@rol_requerido(['Administrador', 'Recepción'])
def abrir_recordatorio_whatsapp(request, id_recordatorio):
    try:
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT destinatario, mensaje
                FROM recordatorios
                WHERE id_recordatorio = %s
                  AND canal = %s
            """, [id_recordatorio, 'whatsapp'])
            recordatorio = cursor.fetchone()

        if not recordatorio:
            return HttpResponse('Recordatorio no encontrado', status=404)

        destinatario, mensaje = recordatorio
        numero_limpio = limpiar_numero_whatsapp(destinatario)
        mensaje_codificado = quote(str(mensaje))
        url_whatsapp = f"https://wa.me/{numero_limpio}?text={mensaje_codificado}"

        with connection.cursor() as cursor:
            cursor.execute("""
                UPDATE recordatorios
                SET estado_envio = %s
                WHERE id_recordatorio = %s
            """, ['enviado', id_recordatorio])

        return redirect(url_whatsapp)
    except Exception as e:
        return HttpResponse(f"Error abriendo WhatsApp: {str(e)}", status=500)


@rol_requerido(['Administrador', 'Recepción'])
def abrir_whatsapp_cita(request, id_cita):
    cita = get_object_or_404(
        Citas.objects.select_related('id_paciente', 'id_estado_cita'),
        id_cita=id_cita
    )

    if not cita.id_paciente or not cita.id_paciente.telefono:
        return HttpResponse('La cita no tiene teléfono de paciente registrado.', status=400)

    mensaje = _mensaje_cita_recordatorio(cita)
    numero_limpio = limpiar_numero_whatsapp(cita.id_paciente.telefono)
    url_whatsapp = f"https://wa.me/{numero_limpio}?text={quote(str(mensaje))}"

    try:
        with connection.cursor() as cursor:
            cursor.execute("""
                INSERT INTO recordatorios (
                    id_cita,
                    canal,
                    tipo_recordatorio,
                    destinatario,
                    mensaje,
                    fecha_programada,
                    estado_envio
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """, [
                cita.id_cita,
                'whatsapp',
                'manual_dashboard',
                cita.id_paciente.telefono,
                mensaje,
                datetime.now(),
                'enviado'
            ])
    except Exception as e:
        print('NO SE PUDO REGISTRAR WHATSAPP MANUAL:', e)

    return redirect(url_whatsapp)


@rol_requerido(['Administrador', 'Recepción'])
def generar_recordatorios_automaticos(request):
    """
    Genera recordatorios manuales para citas de hoy y mañana.
    Se crean como pendientes desde ahora, para que los botones de correo/WhatsApp los encuentren de inmediato.
    """
    creados = 0
    omitidos = 0
    errores = 0

    try:
        citas = _citas_recordatorio_qs()
        fecha_programada = datetime.now()

        for cita in citas:
            try:
                if not cita.id_paciente:
                    omitidos += 1
                    continue

                mensaje = _mensaje_cita_recordatorio(cita)
                canales = [
                    ('correo', 'manual_dashboard', cita.id_paciente.correo),
                    ('whatsapp', 'manual_dashboard', cita.id_paciente.telefono),
                ]

                # Agregar recordatorio al doctor si tiene datos de contacto
                if cita.id_doctor and cita.id_doctor.id_usuario:
                    msg_doctor = _mensaje_cita_recordatorio_doctor(cita)
                    correo_doctor = cita.id_doctor.id_usuario.correo
                    tel_doctor = cita.id_doctor.id_usuario.telefono
                    if correo_doctor:
                        canales.append(('correo', 'manual_dashboard_doctor', correo_doctor))
                    if tel_doctor:
                        canales.append(('whatsapp', 'manual_dashboard_doctor', tel_doctor))
                    # Reemplazar mensaje en canales del doctor
                    canales_doctor_indices = [i for i, c in enumerate(canales) if c[1] == 'manual_dashboard_doctor']
                    for idx in canales_doctor_indices:
                        canal, tipo, dest = canales[idx]
                        canales[idx] = (canal, tipo, dest, msg_doctor)

                # Normalizar canales a tuplas de 4 elementos
                canales = [c if len(c) == 4 else (c[0], c[1], c[2], mensaje) for c in canales]

                for canal, tipo, destinatario, msg in canales:
                    if not destinatario:
                        omitidos += 1
                        continue

                    with connection.cursor() as cursor:
                        cursor.execute("""
                            SELECT COUNT(*)
                            FROM recordatorios
                            WHERE id_cita = %s
                              AND canal = %s
                              AND tipo_recordatorio = %s
                              AND estado_envio = %s
                        """, [cita.id_cita, canal, tipo, 'pendiente'])
                        existe = cursor.fetchone()[0]

                    if existe:
                        omitidos += 1
                        continue

                    with connection.cursor() as cursor:
                        cursor.execute("""
                            INSERT INTO recordatorios (
                                id_cita,
                                canal,
                                tipo_recordatorio,
                                destinatario,
                                mensaje,
                                fecha_programada,
                                estado_envio
                            )
                            VALUES (%s, %s, %s, %s, %s, %s, %s)
                        """, [
                            cita.id_cita,
                            canal,
                            tipo,
                            destinatario,
                            msg,
                            fecha_programada,
                            'pendiente'
                        ])

                    creados += 1

            except Exception as e:
                errores += 1
                print('ERROR GENERANDO RECORDATORIO:', e)

        return HttpResponse(
            f"Recordatorios generados. Creados: {creados}. Omitidos: {omitidos}. Errores: {errores}."
        )

    except Exception as e:
        return HttpResponse(f"Error al generar recordatorios: {str(e)}", status=500)



def api_dashboard_recordatorios(request):
    """
    Dashboard corregido para producción en Render/PostgreSQL.
    No usa funciones MySQL como CURDATE(), DATE_ADD(), MONTH() o YEAR().
    Además evita error 500: si algo falla, devuelve el error en el JSON para poder diagnosticarlo.
    """
    try:
        from decimal import Decimal
        from zoneinfo import ZoneInfo

        def to_float(value):
            if value is None:
                return 0.0
            if isinstance(value, Decimal):
                return float(value)
            try:
                return float(value)
            except Exception:
                return 0.0

        def estado_no_cancelado(qs):
            return qs.exclude(id_estado_cita__nombre_estado__iexact='Cancelada').exclude(
                id_estado_cita__nombre_estado__iexact='Cancelado'
            )

        def nombre_paciente(paciente):
            if not paciente:
                return ''
            return f"{paciente.nombres or ''} {paciente.apellidos or ''}".strip()

        def nombre_doctor(doctor):
            if not doctor or not doctor.id_usuario:
                return ''
            return f"{doctor.id_usuario.nombres or ''} {doctor.id_usuario.apellidos or ''}".strip()

        try:
            ahora_gt = datetime.now(ZoneInfo("America/Guatemala"))
        except Exception:
            ahora_gt = datetime.now()
        hoy = ahora_gt.date()
        hace_7_dias_dt = datetime.now() - timedelta(days=7)
        fin_7_dias = hoy + timedelta(days=7)

        citas_base = Citas.objects.select_related(
            'id_paciente',
            'id_doctor__id_usuario',
            'id_estado_cita'
        )

        citas_validas = estado_no_cancelado(citas_base)

        pagos_base = Pagos.objects.select_related(
            'id_paciente',
            'id_cita',
            'id_tipo_pago'
        )

        pagos_mes_qs = pagos_base.filter(
            fecha_pago__isnull=False,
            fecha_pago__year=hoy.year,
            fecha_pago__month=hoy.month
        )

        ingresos_semanales = pagos_base.filter(
            fecha_pago__isnull=False,
            fecha_pago__gte=hace_7_dias_dt
        ).aggregate(total=Sum('monto'))['total'] or 0

        ingresos_mes = pagos_mes_qs.aggregate(total=Sum('monto'))['total'] or 0

        ultimo_pago_obj = pagos_base.order_by('-fecha_pago', '-id_pago').first()
        ultimo_pago = to_float(ultimo_pago_obj.monto) if ultimo_pago_obj else 0

        total_citas_mes = citas_base.filter(
            fecha_cita__isnull=False,
            fecha_cita__year=hoy.year,
            fecha_cita__month=hoy.month
        ).count()

        canceladas_mes = citas_base.filter(
            fecha_cita__isnull=False,
            fecha_cita__year=hoy.year,
            fecha_cita__month=hoy.month,
            id_estado_cita__nombre_estado__iexact='Cancelada'
        ).count() + citas_base.filter(
            fecha_cita__isnull=False,
            fecha_cita__year=hoy.year,
            fecha_cita__month=hoy.month,
            id_estado_cita__nombre_estado__iexact='Cancelado'
        ).count()

        tasa_cancelaciones = round((canceladas_mes / total_citas_mes) * 100, 2) if total_citas_mes else 0

        citas_con_pago_ids = list(
            Pagos.objects.exclude(id_cita__isnull=True).values_list('id_cita_id', flat=True)
        )

        citas_pasadas_sin_pago = estado_no_cancelado(
            citas_base.filter(fecha_cita__lt=hoy)
        ).exclude(id_cita__in=citas_con_pago_ids)

        data = {
            "ok": True,
            "citas_hoy": citas_validas.filter(fecha_cita=hoy).count(),
            "citas_proximos_7_dias": citas_validas.filter(fecha_cita__gte=hoy, fecha_cita__lte=fin_7_dias).count(),
            "pacientes_con_deuda": citas_pasadas_sin_pago.values('id_paciente').distinct().count(),
            "ingresos_semanales": to_float(ingresos_semanales),
            "ingresos_mes": to_float(ingresos_mes),
            "pagos_mes": pagos_mes_qs.count(),
            "ultimo_pago": ultimo_pago,
            "tasa_cancelaciones": tasa_cancelaciones,
            "canceladas_mes": canceladas_mes,
            "total_citas_mes": total_citas_mes,
            "correos_pendientes": 0,
            "whatsapp_pendientes": 0,
            "total_recordatorios_pendientes": 0,
            "recordatorios_pendientes": [],
            "pagos_recientes": [],
            "citas_proximas": [],
            "pacientes_pendientes_pago": [],
            "citas_estado_mes": [],
            "ultima_actualizacion_servidor": ahora_gt.strftime('%H:%M:%S'),
            "debug_total_citas": Citas.objects.count(),
            "debug_total_pagos": Pagos.objects.count(),
            "debug_hoy_gt": str(hoy),
        }

        for p in pagos_mes_qs.order_by('-fecha_pago', '-id_pago')[:8]:
            data["pagos_recientes"].append({
                "id_pago": p.id_pago,
                "paciente": nombre_paciente(p.id_paciente),
                "monto": to_float(p.monto),
                "fecha_pago": p.fecha_pago.strftime('%Y-%m-%d %H:%M') if p.fecha_pago else '',
                "tipo_pago": p.id_tipo_pago.nombre_tipo_pago if p.id_tipo_pago else 'Sin tipo'
            })

        for c in citas_validas.filter(fecha_cita__gte=hoy, fecha_cita__lte=fin_7_dias).order_by('fecha_cita', 'hora_inicio')[:8]:
            data["citas_proximas"].append({
                "id_cita": c.id_cita,
                "paciente": nombre_paciente(c.id_paciente),
                "doctor": nombre_doctor(c.id_doctor),
                "fecha_cita": str(c.fecha_cita) if c.fecha_cita else '',
                "hora_inicio": str(c.hora_inicio)[:5] if c.hora_inicio else '',
                "estado": c.id_estado_cita.nombre_estado if c.id_estado_cita else 'Sin estado'
            })

        for c in citas_pasadas_sin_pago.order_by('-fecha_cita', '-hora_inicio')[:8]:
            data["pacientes_pendientes_pago"].append({
                "id_cita": c.id_cita,
                "paciente": nombre_paciente(c.id_paciente),
                "fecha_cita": str(c.fecha_cita) if c.fecha_cita else '',
                "hora_inicio": str(c.hora_inicio)[:5] if c.hora_inicio else '',
                "estado": c.id_estado_cita.nombre_estado if c.id_estado_cita else 'Sin estado'
            })

        estados_mes = citas_base.filter(
            fecha_cita__isnull=False,
            fecha_cita__year=hoy.year,
            fecha_cita__month=hoy.month
        ).values('id_estado_cita__nombre_estado').annotate(total=Count('id_cita')).order_by('-total')

        for item in estados_mes:
            data["citas_estado_mes"].append({
                "estado": item.get('id_estado_cita__nombre_estado') or 'Sin estado',
                "total": int(item.get('total') or 0)
            })

        try:
            with connection.cursor() as cursor:
                cursor.execute("""
                    SELECT COUNT(*) FROM recordatorios
                    WHERE canal = %s AND estado_envio = %s
                    AND fecha_programada >= CURRENT_DATE
                """, ['correo', 'pendiente'])
                data["correos_pendientes"] = int(cursor.fetchone()[0] or 0)

                cursor.execute("""
                    SELECT COUNT(*) FROM recordatorios
                    WHERE canal = %s AND estado_envio = %s
                    AND fecha_programada >= CURRENT_DATE
                """, ['whatsapp', 'pendiente'])
                data["whatsapp_pendientes"] = int(cursor.fetchone()[0] or 0)

                cursor.execute("""
                    SELECT id_recordatorio, canal, tipo_recordatorio, destinatario, mensaje, fecha_programada, estado_envio
                    FROM recordatorios
                    WHERE estado_envio = %s
                    AND fecha_programada >= CURRENT_DATE
                    ORDER BY fecha_programada ASC
                    LIMIT 8
                """, ['pendiente'])
                filas = cursor.fetchall()

            data["total_recordatorios_pendientes"] = data["correos_pendientes"] + data["whatsapp_pendientes"]

            for fila in filas:
                data["recordatorios_pendientes"].append({
                    "id_recordatorio": fila[0],
                    "canal": fila[1],
                    "tipo_recordatorio": fila[2],
                    "destinatario": fila[3],
                    "mensaje": fila[4],
                    "fecha_programada": str(fila[5]) if fila[5] else '',
                    "estado_envio": fila[6],
                })
        except Exception as e:
            try:
                connection.rollback()
            except Exception:
                pass
            data["debug_error_recordatorios"] = str(e)

        return JsonResponse(data, json_dumps_params={'ensure_ascii': False})

    except Exception as e:
        import traceback
        return JsonResponse({
            "ok": False,
            "error": str(e),
            "traceback": traceback.format_exc(),
            "citas_hoy": 0,
            "citas_proximos_7_dias": 0,
            "pacientes_con_deuda": 0,
            "ingresos_semanales": 0,
            "ingresos_mes": 0,
            "pagos_mes": 0,
            "ultimo_pago": 0,
            "tasa_cancelaciones": 0,
            "canceladas_mes": 0,
            "total_citas_mes": 0,
            "correos_pendientes": 0,
            "whatsapp_pendientes": 0,
            "total_recordatorios_pendientes": 0,
            "recordatorios_pendientes": [],
            "pagos_recientes": [],
            "citas_proximas": [],
            "pacientes_pendientes_pago": [],
            "citas_estado_mes": [],
            "ultima_actualizacion_servidor": datetime.now().strftime('%H:%M:%S')
        }, json_dumps_params={'ensure_ascii': False})
